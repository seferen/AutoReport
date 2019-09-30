import requests
from docx import Document
from docx.shared import Cm
from _datetime import datetime
import json
import tempfile


class Report:

    def __init__(self, fileName = input('Введите путь до json файла с настройками:\n')):
        self.property = None
        self.TOKEN = None
        self.dashboardName = None
        self.host = None
        self.port = None
        self.graficsDashboards = None
        self.timesFromGrathics = None
        self.loadJsonFile(fileName)


    def parcingDate(self, timeToParse: str):
        return str(int(datetime.strptime(timeToParse, '%d.%m.%Y %H:%M').timestamp() * 1000))


    def getUrl(self, orgId: str, panelId: str, timeStart: str, timeEnd: str):
        return 'http://' + self.host + ':' + self.port + '/render/dashboard-solo/' + self.dashboardName + \
               '?panelId=' + panelId + \
               '&orgId=' + orgId + \
               '&from=' + timeStart + \
               '&to=' + timeEnd + \
               '&theme=light&width=1000&height=500&tz=UTC%2B03%3A00'


    def loadJsonFile(self, fileName):
        try:
            with open(fileName, 'r', encoding='utf8') as f:
                property = json.load(f)
        except Exception:
            print('что то пошло не так, проверьте правильности введенного имени файла!!!')
        try:
            self.TOKEN = property['token']
            self.dashboardName = property['dashboardName']
            self.host = property['host']
            self.port = property['port']
            # Первое значение ogrID, второе значение panelId, третье значение добавляемое наименование таблицы
            self.graficsDashboards = property['graficsDashboards']
            self.timesFromGrathics = property['timesFromGrafics']
        except Exception:
            print("Неверно указаны параметры в json файле ")

    def createAndSaveDocuments(self):

        for time in self.timesFromGrathics:
            datetimeStart: str = self.parcingDate(time[0])
            datetimeEnd: str = self.parcingDate(time[1])

            document = Document()
            document.add_heading('Протокол запуска теста \n' +
                                 time[2], 1)
            document.add_paragraph('Время начала теста:' + time[0], 'Normal')
            document.add_paragraph('Время окончания теста: ' + time[1], 'Normal')
            document.add_paragraph('Продолжительность теста: ', 'Normal')
            document.add_heading('Описание:', 4)
            document.add_paragraph('', 'Normal')
            document.add_heading('Выводы:', 4)
            document.add_paragraph('', 'Normal')
            document.add_heading('Графики:', 3)


            for i in self.graficsDashboards:

                url = self.getUrl(i[0], i[1], datetimeStart, datetimeEnd)
                response = requests.get(url, headers={'Authorization': 'Bearer ' + self.TOKEN})
                if response.status_code == 200:
                    f = tempfile.TemporaryFile()
                    f.write(response.content)
                    f.seek(0)
                    document.add_picture(f, width=Cm(15))
                    f.close()
                    document.add_paragraph(i[2], 'Caption')
                    print('график добавлен: ' + i[2])
                else:
                    print("Чтот то пошло не так с графиком: " + i[2] + "\nresponse code: " + str(response.status_code))

            try:
                document.save(time[2] + '.docx')
            except Exception:
                document.save(time[2] + 'Und' + '.docx')
                pass

            print('done')


if __name__ == '__main__':

    result = Report()

    result.createAndSaveDocuments()
